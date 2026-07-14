import docx

PLACEHOLDER_PATTERNS = ["{{QUESTIONS}}", "[QUESTIONS]", "{QUESTIONS}", "<<QUESTIONS>>"]


def _find_placeholder_paragraph(doc):
    for para in doc.paragraphs:
        text = para.text.strip()
        for pattern in PLACEHOLDER_PATTERNS:
            if pattern.lower() in text.lower():
                return para
    return None


def has_placeholder(template_path: str) -> bool:
    doc = docx.Document(template_path)
    return _find_placeholder_paragraph(doc) is not None


def build_docx_from_template(
    template_path: str,
    output_path: str,
    sections: list,  # [{"name": str, "questions": [{"question":.., "answer":..}]}]
    include_answers: bool = False,
):
    """Opens the college's own .docx template and inserts generated questions
    into it, preserving whatever header/logo/footer/styling the template
    already has. If the template contains a {{QUESTIONS}} marker, questions
    are inserted exactly there. Otherwise they're appended to the end."""
    doc = docx.Document(template_path)
    placeholder = _find_placeholder_paragraph(doc)

    if placeholder is not None:
        placeholder.text = ""
        anchor = placeholder._p
    else:
        spacer = doc.add_paragraph("")
        anchor = spacer._p

    def insert_after(anchor_elem, text, bold=False, indent=False):
        new_para = doc.add_paragraph()
        run = new_para.add_run(text)
        run.bold = bold
        if indent:
            new_para.paragraph_format.left_indent = docx.shared.Inches(0.25)
        anchor_elem.addnext(new_para._p)
        return new_para._p

    current = anchor
    for section in sections:
        current = insert_after(current, section["name"], bold=True)
        for i, q in enumerate(section["questions"], start=1):
            current = insert_after(current, f"Q{i}. {q['question']}")

    if include_answers:
        current = insert_after(current, "")
        current = insert_after(current, "Answer Key", bold=True)
        for section in sections:
            current = insert_after(current, section["name"], bold=True)
            for i, q in enumerate(section["questions"], start=1):
                ans = q.get("answer") or "(no model answer generated)"
                current = insert_after(current, f"A{i}. {ans}", indent=True)

    doc.save(output_path)
    return output_path
